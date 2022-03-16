"""Directions to use:

0. Install the python-docx package with pip. Microsoft Word install is not
   necessary.
1. Revise the name and instruments according to your bingo needs.
2. Run the script.
3. A .docx file with a class set of random bingo boards will be created in your
   current working directory, with the name bingo_board.docx
"""

from random import shuffle

from docx import Document


# change me
BINGO_GAME_NAME = "Musical Instrument Bingo"


# change me
INSTRUMENTS = [
    "electric guitar",
    "acoustic guitar",
    "piano",
    "electric bass",
    "drum set",
    "ukulele",
    "piccolo",
    "flute",
    "clarinet",
    "bass clarinet",
    "soprano saxophone",
    "alto saxophone",
    "tenor saxophone",
    "baritone saxophone",
    "trumpet",
    "slide trumpet",
    "trombone",
    "euphonium",
    "tuba",
    "violin",
    "viola",
    "cello",
    "string bass",
    "timpani",
    "cowbell",
    "congas",
    "timbales",
    "maracas",
    "guiro",
    "clave",
    "triangle",
    "tambourine",
    "theremin",
    "marimba",
    "xylophone",
    "vibraphone",
]


# ---
# script; only change the below if you dare!
# ---


def is_bingo(x, y):
    """Cells at the center of the board are "free" cells"""
    return x in range(2, 4) and y in range(2, 4)


def add_bingo_board(doc):
    """Add a single bingo board table to the given document."""
    table = doc.add_table(rows=6, cols=6)
    table.style = "TableGrid"
    for i, instrument in enumerate(INSTRUMENTS):
        x = i % 6
        y = i // 6
        if is_bingo(x, y):
            table.rows[i % 6].cells[i // 6].text = "FREE"
        else:
            table.rows[i % 6].cells[i // 6].text = instrument


def main():
    """Create a bingo printout for one class, with 25 unique boards."""
    document = Document()
    for _ in range(26):
        shuffle(INSTRUMENTS)
        document.add_heading("Instrument Bingo")

        # add two boards to each page, to allow for a second round
        document.add_paragraph("Round One")
        add_bingo_board(document)
        document.add_paragraph("Round Two")
        add_bingo_board(document)

        document.add_page_break()

    document.save("bingo_board.docx")


if __name__ == "__main__":
    main()
